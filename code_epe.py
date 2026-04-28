import streamlit as st
import mysql.connector
import pandas as pd
import numpy as np
from sklearn.preprocessing import MinMaxScaler
from sklearn.neighbors import NearestNeighbors

# Add these imports
try:
    from docx import Document
    from docx.shared import Inches
    import io  # for in-memory file handling
    WORD_EXPORT_AVAILABLE = True
except ImportError:
    WORD_EXPORT_AVAILABLE = False

def Marks_10_12_Diploma(ModeofEducation):
    if ModeofEducation == "Higher Secondary (or) 12":
        c10 = st.columns(2)
        with c10[0]:
            board10 = st.selectbox(
            "Board of study (10th):",[
            "Tamil Nadu Board of Secondary Education",
            "Central Board of Secondary Education",
            "Indian Certificate of Secondary Education",
            "National Institute of Open Schooling"
            ], key="board10")
        with c10[1]:
            marks_10 = st.number_input("Enter your 10th marks (Percentage %):", value=60.0, min_value=0.0, max_value=100.0, step=1.0)

        c12 = st.columns(2)
        with c12[0]:
            board12 = st.selectbox(
            "Board of study (12th):",[
            "Tamil Nadu Board of Secondary Education",
            "Central Board of Secondary Education",
            "Indian Certificate of Secondary Education",
            "National Institute of Open Schooling"
            ], key="board12_12_diploma")
        with c12[1]:
            Course_12 = st.selectbox(
            "Choose your Stream (12th):",[
            "Computer Science with Maths",
            "Computer Science with Commerce",
            "Biology Science with Maths",
            "Commerce",
            "Arts"
            ], key="Course_12_12_diploma")
        marks_12 = st.number_input("Enter your 12th marks (Percentage %):", value=60.0, min_value=0.0, max_value=100.0, step=1.0)

        if marks_10 >= 34 and marks_12 >= 34:
            dic = {"10": marks_10, "12": marks_12}
            return dic, [Course_12]
        else:
            return 0, ""

    elif ModeofEducation == "Diploma":
        c10 = st.columns(2)
        with c10[0]:
            board10 = st.selectbox(
            "Board of study (10th):",[
            "Tamil Nadu Board of Secondary Education",
            "Central Board of Secondary Education",
            "Indian Certificate of Secondary Education",
            "National Institute of Open Schooling"
            ], key="board10_diploma")
        with c10[1]:
            marks_10 = st.number_input("Enter your 10th marks (Percentage %):", value=60.0, min_value=0.0, max_value=100.0, step=1.0)

        cdip = st.columns(2)
        with cdip[0]:
            Course_Diploma = st.selectbox(
                "Choose your Stream (Diploma):",[
                "Diploma in Engineering",
                "Diploma in Architecture",
                "Diploma in Pharmacy",
                "Diploma in Hotel Management",
                "Diploma in Fashion Design",
                "Diploma in Journalism and Mass Communication",
                "Diploma in Animation and Multimedia",
                "Diploma in Nursing"
                ], key="Course_Diploma")
        with cdip[1]:    
            marks_Diploma = st.number_input("Enter your Diploma marks (Percentage %):", value=60.0, min_value=0.0, max_value=100.0, step=1.0)

        if marks_10 >= 34 and marks_Diploma >= 40:
            dic = {"10": marks_10, "Diploma": marks_Diploma}
            return dic, [Course_Diploma]
        else:
            return 0, ""

    elif ModeofEducation == "12 + Diploma":
        c10 = st.columns(2)
        with c10[0]:
            board10 = st.selectbox(
            "Board of study (10th):",[
            "Tamil Nadu Board of Secondary Education",
            "Central Board of Secondary Education",
            "Indian Certificate of Secondary Education",
            "National Institute of Open Schooling"
            ], key="board10_12_diploma")
        with c10[1]:
            marks_10 = st.number_input("Enter your 10th marks (Percentage %):", value=60.0, min_value=0.0, max_value=100.0, step=1.0)

        c12 = st.columns(2)
        with c12[0]:
            board12 = st.selectbox(
            "Board of study (12th):",[
            "Tamil Nadu Board of Secondary Education",
            "Central Board of Secondary Education",
            "Indian Certificate of Secondary Education",
            "National Institute of Open Schooling"
            ], key="board12_12_diploma")
        with c12[1]:
            Course_12 = st.selectbox(
            "Choose your Stream (12th):",[
            "Computer Science with Maths",
            "Computer Science with Commerce",
            "Biology Science with Maths",
            "Commerce",
            "Arts"
            ], key="Course_12_12_diploma")
        marks_12 = st.number_input("Enter your 12th marks (Percentage %):", value=60.0, min_value=0.0, max_value=100.0, step=1.0)

        cdip = st.columns(2)
        with cdip[0]:
            Course_Diploma = st.selectbox(
                "Choose your Stream (Diploma):",[
                "Diploma in Engineering",
                "Diploma in Architecture",
                "Diploma in Pharmacy",
                "Diploma in Hotel Management",
                "Diploma in Fashion Design",
                "Diploma in Journalism and Mass Communication",
                "Diploma in Animation and Multimedia",
                "Diploma in Nursing"
                ], key="Course_Diploma_12_diploma")
        with cdip[1]:    
            marks_Diploma = st.number_input("Enter your Diploma marks (Percentage %):", value=60.0, min_value=0.0, max_value=100.0, step=1.0)

        if (marks_10 >= 34 and marks_12 >= 34) and marks_Diploma >= 40:
            dic = {"10": marks_10, "12": marks_12, "Diploma": marks_Diploma}
            return dic, [Course_12,Course_Diploma]
        else:
            return 0, ""

def degree_to_course(deg):
    db_config = {
        "host": "localhost",
        "user": "root",
        "password": "sql@123",
        "database": "COLLEGE_DATA"
    }
    conn = mysql.connector.connect(**db_config)
    cursor = conn.cursor()
    
    insert_query = """SELECT distinct course from TN_CLG where degree = %s"""
    cursor.execute(insert_query,(deg,))
    courses = cursor.fetchall()
    c = []
    for i in courses:
        c.append(i[0])
    conn.commit()
    if len(c) == 1:
        return ["-- Select --"] + c
    else:
        return ["-- Select --"] + ["Select all Courses"] + c
    

def course_to_spcl(deg,cour):
    db_config = {
        "host": "localhost",
        "user": "root",
        "password": "sql@123",
        "database": "COLLEGE_DATA"
    }
    conn = mysql.connector.connect(**db_config)
    cursor = conn.cursor()
    
    insert_query = """SELECT distinct specialization from TN_CLG where degree = %s and course = %s"""
    cursor.execute(insert_query,[deg,cour])
    spcl = cursor.fetchall()
    s = []
    for i in spcl:
        s.append(i[0])
    conn.commit()
    if len(s) == 1:
        return ["-- Select --"] + s
    else:
        return ["-- Select --"] + ["Select all Specialization"] + s

def Fee_split(f):
    m = ""
    for i in f:
        if i.isdigit():
            m+=i
    return int(m)

def ML_Mapping(colleges,preference):
    pfr = [1,preference[2],100,0]
    print("Dafault inputs",pfr)
    # location, fee, placement, NIRF
    
    ml = []
    
    for i in colleges:
        marks = st.session_state.marks

        print(i)

        if marks["10"]>float(i[10]):
            Qualify_10 = "True"
        else:
            Qualify_10 = "False"

        if "12" in marks:
            if marks["12"]>float(i[11]):
                Qualify_12 = "True"
            else:
                Qualify_12 = "False"
        else:
            Qualify_12 = "True"
    
        if Qualify_12 == "True" and Qualify_10 == "True":
            if i[9] == preference[0]:l = 1
            elif i[9] == preference[1]:l = 2
            else: l = 3
            f = Fee_split(i[5])
            e = int(i[12])
            n = int(i[8])
            ml.append([i[0],l,f,e,n])
    
    print("Colleges Preference",ml[:10])

    if len(ml) == 0:
        clgs = []

    else:
        df = pd.DataFrame(ml, columns=["ID", "Location", "Fees", "Placement", "NIRF"])

        # Features to use for ranking
        features = ["Location", "Fees", "Placement", "NIRF"]

        # Assign custom importance (lower = more important)
        importance = {
            "Location": 0.5, 
            "Fees": 1,       
            "Placement": 1.5,
            "NIRF": 0.75     
            \
        }

        # Apply importance: Divide by importance values
        for feature, weight in importance.items():  # FIXED VARIABLE NAME
            df[feature] /= weight # Dividing makes the feature smaller → more impact in KNN

        # Normalize the features
        scaler = MinMaxScaler()
        df_scaled = scaler.fit_transform(df[features])

        df_normalized = pd.DataFrame(df_scaled, columns=features)

        print(df_normalized.head(20))

        # Train KNN Model
        knn = NearestNeighbors(n_neighbors=len(ml), metric="manhattan")
        knn.fit(df_scaled)

        # Example Query Point (New Input)
        query_point = np.array([pfr])  # Replace with actual values

        # Apply importance to query point before scaling
        query_point_weighted = np.array([
            query_point[0][0] / importance["Location"],
            query_point[0][1] / importance["Fees"],
            query_point[0][2] / importance["Placement"],
            query_point[0][3] / importance["NIRF"]
        ]).reshape(1, -1)

        query_scaled = scaler.transform(query_point_weighted)

        # Find nearest neighbors
        distances, indices = knn.kneighbors(query_scaled)

        # Ranked DataFrame based on proximity
        ranked_df = df.iloc[indices[0]].copy()
        ranked_df["Distance"] = distances[0]  # Add distance column

        # Sort by distance (lower is better)
        ranked_df = ranked_df.sort_values(by="Distance")

        output = ranked_df["ID"]

        clgs = []
        for i in output:
            for j in colleges:
                if i == j[0]:
                    clgs.append(j)

    return clgs

def print_colleges(clgs,colleges):
    if len(clgs) == 0:
        st.markdown("#### You are not eligible for this course requirements")
                        
    else:
        st.markdown(f"##### 🎓 {len(colleges)} results found. You are eligible for {len(clgs)}")
        st.markdown("---")

        for entry in clgs:
            _, college, degree, course, specialization, fee, subjects, exams, rank, location, _, _, emp = entry

            st.markdown(f"### {college}")  # No need for HTML
            st.markdown(f"##### {degree}  {course} ({specialization})")
            col1, col2 = st.columns([1, 2])  # Creates two columns for better layout

            with col1:
                st.write(f"**NIRF Rank:** {rank}")
                st.write(f"**Placement Rate:** {emp}%")
                st.write(f"**Location:** {location}")

            with col2:
                st.write(f"**Fee Structure:** {fee}")
                st.write(f"**Subjects Required:** {subjects}")
                st.write(f"**Competitive Exams:** {exams}")

            st.markdown("---")

    return clgs # return the list of colleges

def Get_College(lst,Preference):
    db_config = {
        "host": "localhost",
        "user": "root",
        "password": "sql@123",
        "database": "COLLEGE_DATA"
    }
    conn = mysql.connector.connect(**db_config)
    cursor = conn.cursor()

    if lst[1] == "":
        insert_query = """SELECT * from TN_CLG where degree = %s"""
        cursor.execute(insert_query,[lst[0]])
    elif lst[2] == "":
        insert_query = """SELECT * from TN_CLG where degree = %s and course = %s"""
        cursor.execute(insert_query,[lst[0],lst[1]])
    else:
        insert_query = """SELECT * from TN_CLG where degree = %s and course = %s and specialization = %s"""
        cursor.execute(insert_query,[lst[0],lst[1],lst[2]])

    colleges = cursor.fetchall()
    conn.commit()

    if len(colleges) == 0:
        st.markdown("#### No Record Available")
        return [] # return empty list if no college found
    else:
        clgs = ML_Mapping(colleges,Preference)
        final_colleges = print_colleges(clgs,colleges)
        return final_colleges  # return list of colleges to be put in word doc.

def Get_College_Dip(c,Preference):  
    db_config = {
        "host": "localhost",
        "user": "root",
        "password": "sql@123",
        "database": "COLLEGE_DATA"
    }
    conn = mysql.connector.connect(**db_config)
    cursor = conn.cursor()

    insert_query = """SELECT * from TN_CLG where course = %s"""
    cursor.execute(insert_query,[c])

    colleges = cursor.fetchall()
    conn.commit()

    if len(colleges) == 0:
        st.markdown("#### No Record Available")
        return []
    else:
        clgs = ML_Mapping(colleges,Preference)
        final_colleges = print_colleges(clgs,colleges)
        return final_colleges
    
def Find_Degrees(x,y):

    Engineering = ['B.E', 'B.Tech', 'Integrated M.Tech', 'Minor Degree Program']
    Arts_and_Management = ['BA', 'BBA', 'B.Com', 'B.Voc', 'BCA', 'Integrated M.Ed', 'BA Hons', 'B.PES', 'B.P.Ed', 'Integrated MA', 'BVA', 'BHMCT', 'B.LIB.I.SC', 'BFA', 'B.Ed', 'B.Lit']
    Medical = ['BPT', 'BHMS', 'B.Pharm', 'D.Pharm', 'Pharm.D', 'BDS', 'B.Optom', 'MBBS', 'P.B.B.Sc'] 
    Healthcare = ['BA LLB Hons', 'B.Com LLB Hons', 'LLB', 'BBA LLB Hons', 'LLB Hons']
    Law = ['BA LLB Hons', 'B.Com LLB Hons', 'LLB', 'BBA LLB Hons', 'LLB Hons']
    Design = ['B.Arch', 'B.Des']
    Science = ['B.Sc', 'B.Sc Hons', 'Integrated M.Sc', 'Integrated MS', 'B.S']

    Stream1 = ["-- Select --"] + Engineering + Arts_and_Management + Healthcare + Law + Design + Science
    Stream2 = ["-- Select --"] + Arts_and_Management + Law
    Stream3 = ["-- Select --"] + Medical + Engineering + Arts_and_Management + Healthcare + Law+ Design + Science
    Stream4 = ["-- Select --"] + Arts_and_Management + Law
    Stream5 = ["-- Select --"] + Arts_and_Management + Law

    degs_12 = []
    degs_dip = []
    if y[0] == "Computer Science with Maths":
        degs_12 = Stream1
    elif y[0] == "Computer Science with Commerce":
        degs_12 = Stream2
    elif y[0] == "Biology Science with Maths":
        degs_12 = Stream3
    elif y[0] == "Commerce":
        degs_12 = Stream4
    elif y[0] == "Arts":
        degs_12 = Stream5
    
    if y[-1] == "Diploma in Engineering":
        degs_dip = ['Aeronautical Engineering', 'Aerospace Engineering', 'Aerospace Engineering and Applied Mechanics', 'Agricultural Engineering', 'Aircraft Maintenance Engineering', 'Automobile Engineering', 'Automotive Engineering (ARAI)', 'Bioengineering', 'Biological Engineering', 'Biomedical Engineering', 'Biomedical Instrumentation Engineering', 'Biotechnology and Biochemical Engineering', 'Chemical and Electrochemical Engineering', 'Chemical Engineering', 'Civil and Structural Engineering', 'Civil Engineering', 'Civil Engineering (Tamil Medium)', 'Civil Engineering and environmental engineering', 'Civil Engineering and Geotechnical engineering', 'Civil Engineering with Computer Applications', 'Computer and Communication Engineering', 'Computer and Engineering', 'Computer Engineering', 'Computer Science and Engineering', 'Computer Science and Engineering (Powered by Virtusa)', 'Computer Science and Engineering (Tamil Medium)', 'Computer Science and Engineering (WIP)', 'Computer Science and Engineering and Business Systems', 'Computer Science and Engineering and Business Systems (in collaboration with TCS)', 'Computer Science and Medical Engineering', 'Electrical and Computer Engineering', 'Electrical and Computer Science Engineering', 'Electrical and Electronics Engineering', 'Electrical and Electronics Engineering(Sandwich)', 'Electrical Engineering', 'Electrical Engineering and Applied Mechanics ', 'Electronics & Instrumentation Engineering', 'Electronics and Communication Engineering', 'Electronics and Communication Engineering ', 'Electronics and Computer Engineering', 'Electronics and Instrumentation Engineering', 'Electronics and Telecommunication Engineering', 'Electronics Engineering', 'Energy and Environmental Engineering', 'Engineering design', 'Engineering Design (Automotive Engineering)', 'Engineering Design (biomedical engineering)', 'Engineering Physics', 'Environmental Engineering', 'Fire Technology And Safety Engineering', 'Food Processing & Engineering', 'Food Processing and Engineering', 'Geo Informatics Engineering', 'Industrial Engineering', 'Industrial Engineering and Management', 'Instrumentation and Control Engineering', 'Manufacturing Engineering', 'Marine Engineering', 'Material Science and Engineering', 'Mechanical and Automation Engineering', 'Mechanical Engineering', 'Mechanical Engineering (Sandwich)', 'Mechanical Engineering (Tamil Medium)', 'Mechanical Engineering and Smart Manufacturing', 'Mechanical Engineering(Sandwich)', 'Mechatronics Engineering', 'Medical Sciences & Engineering', 'Metallurgical and Materials Engineering', 'Metallurgical Engineering', 'Naval Architecture and Ocean Engineering', 'Naval Architecture and Offshore Engineering', 'Ocean Engineering andApplied Mechanics ', 'Petrochemical Engineering', 'Petroleum Engineering', 'Pharmaceutical Engineering', 'Polymer Engineering', 'Production Engineering', 'Production Engineering(Sandwich)', 'Robotics and Automation Engineering', 'Safety and Fire Engineering', 'Software Engineering', 'Sustainability Engineering', 'Transportation Engineering']
    elif y[-1] == "Diploma in Architecture":
        degs_dip = ["Architecture"]
    elif y[-1] == "Diploma in Pharmacy":
        degs_dip = ["Pharmacy"]
    elif y[-1] == "Diploma in Hotel Management":
        degs_dip = ['Hotel Management and Catering Technology', 'Hotel Management & Tourism', 'Hospitality and Hotel Administration']
    elif y[-1] == "Diploma in Fashion Design":
        degs_dip = ['Costume Design and Fashion', 'Fashion Technology', 'Fashion Design', 'Fashion and Apparel Design']
    elif y[-1] == "Diploma in Journalism and Mass Communication":
        degs_dip = ['Journalism', 'Digital Journalism', 'Journalism and Mass Communication', 'Journalism and Elecronic Media', 'Journalism and Visual Communications']
    elif y[-1] == "Diploma in Animation and Multimedia":
        degs_dip = ['Gaming Design', 'Game Art Design','Graphics and Creative Designing','Editing and Sound Design','Visual Arts', 'Visual Communication', 'Animation and Visual Effects', 'Animation and Visual Communication', 'Journalism and Visual Communications']
    elif y[-1] == "Diploma in Interior Design":
        degs_dip = ["Interior Design","Interior Design and Resource Management"]  
    elif y[-1] == "Diploma in Nursing":
        degs_dip = ['Nursing', 'Nursing (Basic)', 'Nursing (Post Basic)']        
    return [degs_12,degs_dip]

def Filter_Preference():
    st.write("Enter you preferred location")
    cols = st.columns(2)
    l = ['Chengalpattu', 'Chennai', 'Coimbatore', 'Cuddalore', 'Kanyakumari', 'Madurai', 'Perambalur', 'Salem', 'Sivagangai', 'Thanjavur', 'Thoothukudi', 'Tiruchirappalli', 'Tirunelveli', 'Tirupattur', 'Tiruvarur', 'Virudhunagar']
    
    with cols[0]:
        l1 = st.selectbox("First choice preference:",l)
    with cols[1]:
        l2 = st.selectbox("Second choice preference:",l)
    
    st.write("Enter your preferred Fee Range")
    Fee = st.slider("Select a value", min_value=1000, max_value=250000, value=100000, step=1000, label_visibility= "hidden")
                    
    return [l1,l2,Fee,0,0]

# Function to create the Word document
def create_word_document(name, colleges):
    document = Document()
    document.add_heading(f"College Suggestions for {name}", 0)

    if not colleges:
        document.add_paragraph("No colleges found based on the provided criteria.")
    else:
        for entry in colleges:
            _, college, degree, course, specialization, fee, subjects, exams, rank, location, _, _, emp = entry

            document.add_heading(college, level=1)
            document.add_paragraph(f"{degree}  {course} ({specialization})")

            # Create a table for better organization
            table = document.add_table(rows=0, cols=2)

            # Add data to the table
            row_cells = table.add_row().cells
            row_cells[0].text = "**NIRF Rank:**"
            row_cells[1].text = str(rank)

            row_cells = table.add_row().cells
            row_cells[0].text = "**Placement Rate:**"
            row_cells[1].text = f"{emp}%"

            row_cells = table.add_row().cells
            row_cells[0].text = "**Location:**"
            row_cells[1].text = location

            row_cells = table.add_row().cells
            row_cells[0].text = "**Fee Structure:**"
            row_cells[1].text = fee

            row_cells = table.add_row().cells
            row_cells[0].text = "**Subjects Required:**"
            row_cells[1].text = subjects

            row_cells = table.add_row().cells
            row_cells[0].text = "**Competitive Exams:**"
            row_cells[1].text = exams
            
            document.add_paragraph()  # Add space between colleges

    # Save the document to in-memory file
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)  # Go to the beginning of the stream before returning

    return doc_io

st.title("🎓 College Finder")

if "form" not in st.session_state:
    st.session_state.form = False

if "marks" not in st.session_state:
    st.session_state.marks = {}

if "deg_list" not in st.session_state:
    st.session_state.deg_list = []
# Get user details
name = st.text_input("Enter your name:")
info = []
if name:
    st.write(f"Hello, {name}! 👋")

    ModeofEducation = st.selectbox(
        "Mention your level of Education:",
        ["-- Select --",
        "Higher Secondary (or) 12", 
        "Diploma", 
        "12 + Diploma"]
        )
    
    if ModeofEducation!="-- Select --":
        with st.form("my_form"):    
            st.write("Fill your Academic Information -")
            x,y = Marks_10_12_Diploma(ModeofEducation)    
            if st.form_submit_button("Find Degrees"):
                st.session_state.form = True
                if x == 0:
                    st.write("Sorry, your marks does not meet the eligibility criteria for an undergraduate program.")      
                else:
                    print(x)
                    st.session_state.marks = x
                    degs = Find_Degrees(x,y)
                    st.session_state.deg_list = degs
        
        if st.session_state.form == True and x!=0:
            degs = st.session_state.deg_list
            final_colleges = []

            if degs[1] == []:
                st.write("You are eligible for Regular Entry")
                degree = st.selectbox("Select your degree:", degs[0])
                if degree!= "-- Select --" :
                    cour = degree_to_course(degree)
                    course = st.selectbox("Select your Course:", cour)

                    if course == "Select all Courses":
                        Preference = Filter_Preference()
                        if st.button("Find Colleges",key = "Find1"):
                            final_colleges = Get_College([degree,"",""],Preference)

                    elif course!="-- Select --":
                        spcl = course_to_spcl(degree,course)
                        specialization = st.selectbox("Select your Specialization:", spcl)

                        if specialization == "Select all Specialization":
                            Preference = Filter_Preference()
                            if st.button("Find Colleges",key = "Find2"):
                                final_colleges = Get_College([degree,course,""],Preference)
                        
                        elif specialization!="-- Select --":
                            Preference = Filter_Preference()
                            if st.button("Find Colleges",key = "Find3"):
                                final_colleges = Get_College([degree,course,specialization],Preference) 
            elif degs[0] == []:
                st.write("You are eligible for Lateral Entry")
                c = ["-- Select --"] + degs[1]
                cdip = st.selectbox("Select your course:", c)
                
                if cdip!= "-- Select --" :
                    Preference = Filter_Preference()
                    if st.button("Find Colleges",key = "1Find1"):
                        final_colleges = Get_College_Dip(cdip,Preference)
            
            else:
                st.write("You are eligible for both Lateral Entry and Regular Entry")
                lat_reg = st.selectbox("Choose your Preference:",["Regular Entry","Lateral Entry"])
                if lat_reg == "Regular Entry":
                    degree = st.selectbox("Select your degree:", degs[0])
                    if degree!= "-- Select --" :
                        cour = degree_to_course(degree)
                        course = st.selectbox("Select your Course:", cour)

                        if course == "Select all Courses":
                            Preference = Filter_Preference()
                            if st.button("Find Colleges",key = "Find1"):
                                final_colleges = Get_College([degree,"",""],Preference)

                        elif course!="-- Select --":
                            spcl = course_to_spcl(degree,course)
                            specialization = st.selectbox("Select your Specialization:", spcl)

                            if specialization == "Select all Specialization":
                                Preference = Filter_Preference()
                                if st.button("Find Colleges",key = "Find2"):
                                    final_colleges = Get_College([degree,course,""],Preference)
                            
                            elif specialization!="-- Select --":
                                Preference = Filter_Preference()
                                if st.button("Find Colleges",key = "Find3"):
                                    final_colleges = Get_College([degree,course,specialization],Preference) 
                elif lat_reg == "Lateral Entry":
                    c = ["-- Select --"] + degs[1]
                    cdip = st.selectbox("Select your course:", c)
                    
                    if cdip!= "-- Select --" :
                        Preference = Filter_Preference()
                        if st.button("Find Colleges",key = "1Find1"):
                            final_colleges = Get_College_Dip(cdip,Preference)
        
            # Generate Word document
            if WORD_EXPORT_AVAILABLE:
                doc_io = create_word_document(name, final_colleges)
                
                # Provide download button for the Word document
                if doc_io:
                    st.download_button(
                        label="Download College Suggestions as Word Document",
                        data=doc_io.getvalue(),
                        file_name=f"college_suggestions_for_{name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("Word export is not available. Install python-docx to enable this feature.")
    else:
        st.warning("Please select your level of education.")
